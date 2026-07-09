"""
docx_parser.py
──────────────
Extracts raw plain text from DOCX files using python-docx.
Preserves paragraph breaks and includes text from tables.
"""

from docx import Document


def parse_docx(path: str) -> str:
    """
    Extract and return all text from a DOCX file at `path`.

    Extracts text from:
    - All body paragraphs (in order)
    - All table cells (row-by-row, cell-by-cell)

    Args:
        path: Absolute or relative path to the DOCX file.

    Returns:
        A single string containing all extracted text, with paragraphs
        separated by newlines.

    Raises:
        FileNotFoundError: If the file does not exist.
        Exception: Re-raises any python-docx parsing errors.
    """
    doc = Document(path)
    parts: list[str] = []

    for block in doc.element.body:
        # Paragraph element
        tag = block.tag.split("}")[-1] if "}" in block.tag else block.tag
        if tag == "p":
            # Find the matching Paragraph object
            for para in doc.paragraphs:
                if para._element is block:
                    text = para.text.strip()
                    if text:
                        parts.append(text)
                    break
        # Table element
        elif tag == "tbl":
            for table in doc.tables:
                if table._element is block:
                    for row in table.rows:
                        row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_cells:
                            parts.append(" | ".join(row_cells))
                    break

    return "\n".join(parts)
